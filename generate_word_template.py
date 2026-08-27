from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NFSU – National Forensic Sciences University\n')
run.font.bold = True
run.font.size = Pt(14)
run = p.add_run('Goa Campus\n')
run.font.size = Pt(12)
run = p.add_run('Near Goa Dairy, Ponda – 403 401, Goa – India.')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COURSE FEES RECEIPT')
run.font.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Receipt No.: ').bold = True
p.add_run('__________________________')
p.add_run('    Date: ').bold = True
p.add_run('__________________________')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Received From Shri./Smt./Kum.: ').bold = True
p.add_run('__________________________________________________________')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Rs.: ').bold = True
p.add_run('__________________________  ')
p.add_run('(in words): ').bold = True
p.add_run('__________________________________________________________________')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('in Cash / by Cheque / D.D. on account of fees as under for ').bold = True
p.add_run('I / II / III / IV ')
p.add_run('semester for the course of ').bold = True
p.add_run('__________________________________________________________')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Name of the School: ').bold = True
p.add_run('__________________________________________________________')

doc.add_paragraph()

rows = [
    ['Sr. No.', 'Particulars', 'Amount Rs.'],
]
items = [
    'Tuition Fees',
    'Library Fee',
    'Caution Money Including Library (Refundable)',
    'Registration Fee',
    'Enrollment Fee',
    'Computer Fees with internet facility (Per Semester)',
    'Semester Examination Fee',
    'Semester end Examination Transcript Fee',
    'Orientation Programme Fee',
    'Identity Card Fee',
    'Laboratory Fee',
    'Eligibility Certificate Fees',
    'Convocation Fee',
    'Student Activity Fee',
]
for i, item in enumerate(items, start=1):
    rows.append([str(i), item, ''])
rows.append(['', 'TOTAL', ''])

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

for col_idx, heading in enumerate(rows[0]):
    cell = table.cell(0, col_idx)
    cell.text = heading
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(11)

for row_data in rows[1:]:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)


doc.add_paragraph()
doc.add_paragraph('Subject to realisation of cheque')
doc.add_paragraph('Out station cheques will not be accepted.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Executive Registrar')
run.font.bold = True
p.add_run(' ' * 20)
run = p.add_run('Cashier')
run.font.bold = True

doc.save('fee_receipt_template.docx')
print('Created fee_receipt_template.docx')
